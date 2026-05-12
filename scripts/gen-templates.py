"""Generate native xlsx + docx templates for /public/templates/.
Run from project root: python scripts/gen-templates.py
"""
import os
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

OUT = "public/templates"
os.makedirs(OUT, exist_ok=True)

# ============================================================
# Helpers
# ============================================================
BRAND = "1F4E79"
BRAND_LIGHT = "DDEBF7"
ACCENT = "C00000"
HEADER_FILL = PatternFill(start_color=BRAND, end_color=BRAND, fill_type="solid")
HEADER_LIGHT = PatternFill(start_color=BRAND_LIGHT, end_color=BRAND_LIGHT, fill_type="solid")
HEADER_FONT = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
BOLD = Font(name="Calibri", size=11, bold=True)
CENTER = Alignment(horizontal="center", vertical="center", wrap_text=True)
LEFT = Alignment(horizontal="left", vertical="center", wrap_text=True)
THIN = Side(border_style="thin", color="666666")
BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)

def xlsx_set_widths(ws, widths):
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

def xlsx_title_row(ws, row, title, ncols, color=BRAND):
    ws.merge_cells(start_row=row, end_row=row, start_column=1, end_column=ncols)
    cell = ws.cell(row=row, column=1, value=title)
    cell.font = Font(name="Calibri", size=16, bold=True, color="FFFFFF")
    cell.alignment = CENTER
    cell.fill = PatternFill(start_color=color, end_color=color, fill_type="solid")
    ws.row_dimensions[row].height = 30

def xlsx_section(ws, row, title, ncols, color=BRAND_LIGHT, dark=False):
    ws.merge_cells(start_row=row, end_row=row, start_column=1, end_column=ncols)
    cell = ws.cell(row=row, column=1, value=title)
    cell.font = Font(name="Calibri", size=12, bold=True, color="FFFFFF" if dark else "1F4E79")
    cell.alignment = LEFT
    cell.fill = PatternFill(start_color=color, end_color=color, fill_type="solid")
    ws.row_dimensions[row].height = 22

def xlsx_header(ws, row, headers):
    for i, h in enumerate(headers, 1):
        c = ws.cell(row=row, column=i, value=h)
        c.font = HEADER_FONT
        c.alignment = CENTER
        c.fill = HEADER_FILL
        c.border = BORDER
    ws.row_dimensions[row].height = 28

def xlsx_data_row(ws, row, values):
    for i, v in enumerate(values, 1):
        c = ws.cell(row=row, column=i, value=v)
        c.alignment = LEFT if i <= 2 else CENTER
        c.border = BORDER
        c.font = Font(name="Calibri", size=10)
    ws.row_dimensions[row].height = 22

def xlsx_brand_footer(ws, row, ncols):
    ws.merge_cells(start_row=row, end_row=row, start_column=1, end_column=ncols)
    cell = ws.cell(row=row, column=1,
                   value="© Atlantis NDT — Free template. Customize for your specific application. "
                         "Demo: info@atlantisndt.com | atlantisndt.com")
    cell.font = Font(name="Calibri", size=9, italic=True, color="666666")
    cell.alignment = CENTER

# ============================================================
# 1) NDT Inspection Checklist (XLSX)
# ============================================================
def make_inspection_checklist():
    wb = Workbook()
    ws = wb.active
    ws.title = "NDT Inspection Checklist"
    xlsx_set_widths(ws, [4, 50, 12, 12, 12, 25])

    xlsx_title_row(ws, 1, "NDT Inspection Checklist — Pre / During / Post", 6)
    ws.cell(row=2, column=1, value="Project:").font = BOLD
    ws.cell(row=2, column=4, value="Date:").font = BOLD
    ws.cell(row=3, column=1, value="Client:").font = BOLD
    ws.cell(row=3, column=4, value="Inspector:").font = BOLD
    ws.cell(row=4, column=1, value="Asset / Equipment:").font = BOLD
    ws.cell(row=4, column=4, value="Method:").font = BOLD
    ws.cell(row=5, column=1, value="Reference Standards:").font = BOLD
    ws.cell(row=5, column=4, value="Procedure No.:").font = BOLD

    sections = [
        ("PRE-INSPECTION (before mobilization)", [
            "Inspection procedure reviewed and approved (latest revision)",
            "Inspector certification current — ASNT/ISO 9712/PCN scope verified",
            "Calibration certificates current — UT block, gauge, source activity",
            "Reference standards available — applicable code (ASME V, API 510/570/653)",
            "Client written practice / SAEP / ACS / DEP reviewed",
            "PPE issued — coveralls, hard hat, safety glasses, gloves, hearing protection",
            "Work permit / PTW issued — hot work / confined space / radiation as applicable",
            "JSA / risk assessment completed and signed by all crew",
            "Equipment functional check completed on-site",
            "Acceptance criteria reviewed and agreed with client / engineer",
            "Drawings / P&ID / isometric / weld map available",
            "Previous inspection records reviewed (if re-inspection)",
        ]),
        ("DURING INSPECTION", [
            "Surface preparation per procedure — cleaning, grinding, etching",
            "Couplant / penetrant / developer applied per procedure",
            "Reference standard checked — block, IIW, V1/V2, ASME basic block",
            "Probe / source / detector verified for the technique",
            "Scan plan executed — coverage % verified",
            "Indications marked + photographed with scale + location reference",
            "Indication sizing per procedure — length / depth / height",
            "Data recorded in field log / report draft",
            "Each indication classified — relevant / non-relevant / acceptable / reject",
            "Hold points notified to client / engineer in advance",
            "Operating-condition data logged (temperature, pressure) where required",
            "Reference standards re-verified at end of shift (calibration check)",
        ]),
        ("POST-INSPECTION", [
            "Inspection report draft completed — code-format (API / ASME)",
            "Photos and data attachments compiled",
            "Acceptance / reject summary clear and unambiguous",
            "Repair / rework instructions issued for rejected items",
            "Report reviewed by Level III / supervisor",
            "Customer-format formatting applied (per client written practice)",
            "Distribution list — owner, inspector body, regulator copies",
            "Records archived per retention schedule (typically 7 years)",
            "Site demobilization — equipment recovered, calibration logged",
            "Hazardous materials disposal — penetrant fluids, developer, radiography waste",
            "Lessons learned captured for next campaign",
            "Invoice / time sheet finalized for billing",
        ]),
    ]

    row = 7
    for section_title, items in sections:
        row += 1
        xlsx_section(ws, row, section_title, 6, color=BRAND, dark=True)
        row += 1
        xlsx_header(ws, row, ["#", "Item", "Pass / Fail / N/A", "Initial", "Date", "Comments / Reference"])
        for i, item in enumerate(items, 1):
            row += 1
            xlsx_data_row(ws, row, [i, item, "", "", "", ""])

    row += 2
    xlsx_section(ws, row, "Sign-off", 6, color=BRAND_LIGHT)
    row += 1
    xlsx_header(ws, row, ["Role", "Name", "ASNT Cert No.", "Method/Level", "Signature", "Date"])
    for role in ["Inspector", "Senior Inspector / Witness", "Level III Review", "Client Representative"]:
        row += 1
        xlsx_data_row(ws, row, [role, "", "", "", "", ""])

    row += 2
    xlsx_brand_footer(ws, row, 6)
    wb.save(f"{OUT}/ndt-inspection-checklist.xlsx")

# ============================================================
# 2) API 653 Tank Inspection Template (XLSX) — multi-sheet
# ============================================================
def make_api653_template():
    wb = Workbook()

    # Sheet: Cover
    cv = wb.active
    cv.title = "Cover"
    xlsx_set_widths(cv, [4, 30, 30, 30])
    xlsx_title_row(cv, 1, "API 653 Aboveground Storage Tank Inspection Record", 4)
    info = [
        ("Tank Tag No.", "Tank Description"),
        ("Site / Location", "Inspection Date"),
        ("Diameter (ft / m)", "Height (ft / m)"),
        ("Product / Service", "Specific Gravity"),
        ("Design Standard", "Year of Construction"),
        ("Last Internal Inspection", "Next Internal Due"),
        ("Last External Inspection", "Next External Due"),
        ("Foundation Type", "Roof Type (Fixed / Floating / IFR)"),
        ("Shell Material", "Floor Material"),
        ("Lining (Y/N + type)", "Cathodic Protection (Y/N + type)"),
        ("Inspector Name", "API 653 Cert No."),
        ("Inspection Type", "Inspection Result (Pass / Conditional / Fail)"),
    ]
    r = 3
    for left, right in info:
        cv.cell(row=r, column=1, value=left).font = BOLD
        cv.cell(row=r, column=2, value="").border = BORDER
        cv.cell(row=r, column=3, value=right).font = BOLD
        cv.cell(row=r, column=4, value="").border = BORDER
        r += 1
    r += 2
    xlsx_brand_footer(cv, r, 4)

    # Sheet: Shell UT Thickness
    sh = wb.create_sheet("Shell UT Thickness")
    xlsx_set_widths(sh, [10, 12, 14, 14, 14, 14, 14, 14, 14, 14, 14, 16])
    xlsx_title_row(sh, 1, "API 653 §6.4 Shell Thickness Measurement", 12)
    xlsx_header(sh, 3, ["Course", "Elev (ft)", "Nominal t (in)", "Min t1 (in)", "Min t2 (in)", "Min t3 (in)", "Min t4 (in)", "Min t5 (in)", "Min t6 (in)", "Avg t (in)", "Min reading", "Remaining Life (yr)"])
    for c in range(1, 11):
        sh.cell(row=4 + c - 1, column=1, value=f"Course {c}")
    for r in range(4, 14):
        for c in range(1, 13):
            sh.cell(row=r, column=c).border = BORDER
            if c > 1:
                sh.cell(row=r, column=c).alignment = CENTER

    # Sheet: Floor / Bottom
    fl = wb.create_sheet("Floor Inspection")
    xlsx_set_widths(fl, [4, 35, 14, 14, 14, 25])
    xlsx_title_row(fl, 1, "API 653 §6.3 Floor / Bottom Inspection", 6)
    xlsx_header(fl, 3, ["#", "Inspection Item", "Pass / Fail", "Min Thickness", "Required Min", "Comments"])
    floor_items = [
        "Sketch plate inspection — visual external",
        "Annular ring thickness measurement",
        "Floor plate UT thickness — center grid (per API 653 §6.3.2)",
        "Floor seam weld visual inspection",
        "Floor scanning — MFL / vacuum box / floor scanner (if internal)",
        "Floor lining condition (if applicable)",
        "Sump / drainage system condition",
        "Underside corrosion observed (Y/N + extent)",
        "Foundation chime contact / settlement measurement",
        "Cathodic protection survey results",
    ]
    for i, item in enumerate(floor_items, 1):
        xlsx_data_row(fl, 3 + i, [i, item, "", "", "", ""])

    # Sheet: Roof Inspection
    rf = wb.create_sheet("Roof Inspection")
    xlsx_set_widths(rf, [4, 35, 14, 14, 14, 25])
    xlsx_title_row(rf, 1, "API 653 §6.5 Roof Inspection", 6)
    xlsx_header(rf, 3, ["#", "Inspection Item", "Pass / Fail", "Min Thickness", "Required Min", "Comments"])
    roof_items = [
        "Fixed roof / IFR external visual",
        "Roof plate UT thickness — grid pattern",
        "Floating roof primary seal condition",
        "Floating roof secondary seal condition",
        "Roof drain operability",
        "Roof support / rafter condition",
        "Roof manway / nozzle reinforcing pad",
        "Roof emergency vent function",
        "Walking lane / hand-rail integrity",
        "Atmospheric venting / PV valve check",
    ]
    for i, item in enumerate(roof_items, 1):
        xlsx_data_row(rf, 3 + i, [i, item, "", "", "", ""])

    # Sheet: Settlement Survey
    ss = wb.create_sheet("Settlement Survey")
    xlsx_set_widths(ss, [8, 16, 16, 16, 20])
    xlsx_title_row(ss, 1, "API 653 Annex B Settlement Survey", 5)
    xlsx_header(ss, 3, ["Station", "Elev (ft / m)", "Diff vs Mean", "% Diff", "Status (Pass/Fail)"])
    for s in range(1, 25):
        xlsx_data_row(ss, 3 + s, [f"S{s}", "", "", "", ""])

    # Sheet: Recommendations
    rc = wb.create_sheet("Recommendations")
    xlsx_set_widths(rc, [4, 14, 50, 14, 14, 20])
    xlsx_title_row(rc, 1, "Recommendations & Follow-up Actions", 6)
    xlsx_header(rc, 3, ["#", "Priority", "Description", "API 653 Ref.", "Due Date", "Responsible"])
    for i in range(1, 11):
        xlsx_data_row(rc, 3 + i, [i, "", "", "", "", ""])

    # Sheet: Sign-off
    so = wb.create_sheet("Sign-off")
    xlsx_set_widths(so, [4, 25, 20, 18, 14])
    xlsx_title_row(so, 1, "Sign-off", 5)
    xlsx_header(so, 3, ["#", "Role", "Name", "API 653 Cert No.", "Date"])
    for i, r in enumerate(["API 653 Inspector", "Owner / Operator Engineer", "Independent Review", "Regulatory (if applicable)"], 1):
        xlsx_data_row(so, 3 + i, [i, r, "", "", ""])

    xlsx_brand_footer(so, 10, 5)
    wb.save(f"{OUT}/api-653-tank-inspection-template.xlsx")

# ============================================================
# 3) NDT Safety Checklist (XLSX)
# ============================================================
def make_safety_checklist():
    wb = Workbook()
    ws = wb.active
    ws.title = "NDT Safety Checklist"
    xlsx_set_widths(ws, [4, 50, 12, 12, 12, 25])

    xlsx_title_row(ws, 1, "NDT Safety Checklist — Worksite, RT, Electrical, Chemical, Confined Space", 6)
    ws.cell(row=2, column=1, value="Project / Site:").font = BOLD
    ws.cell(row=2, column=4, value="Date:").font = BOLD
    ws.cell(row=3, column=1, value="Crew Lead:").font = BOLD
    ws.cell(row=3, column=4, value="JSA Ref.:").font = BOLD

    sections = [
        ("GENERAL WORKSITE", [
            "Site induction / orientation completed for all crew",
            "PTW (Permit to Work) issued and posted at worksite",
            "Hazard communication briefing complete — site-specific risks reviewed",
            "Emergency assembly point identified — route and muster station",
            "First-aid kit on-site + AED location known",
            "Fire extinguishers on-site, current inspection tag",
            "Communication — radio / mobile signal verified",
            "PPE — hard hat, safety glasses, gloves, FRC, steel-toe boots",
            "Lockout / Tagout (LOTO) verified where applicable",
            "Spill kit on-site for chemical / fluid containment",
            "Site speed limit + traffic rules observed",
            "Buddy system in place for elevated / isolated work",
        ]),
        ("RADIATION SAFETY (RT)", [
            "Radiography work permit issued — DOT / state authority compliance",
            "Source serial number + activity verified at SOE start",
            "Source / collimator transport documentation per IAEA TS-R-1",
            "Survey meter calibrated within 12 months — battery checked",
            "Survey meter functional check at start of shift — known source",
            "Dosimetry — TLD / OSL badge + pocket dosimeter for all radiographer crew",
            "Radiation area barricades + signage installed (high / low / restricted)",
            "Boundary radiation level confirmed — 2 mR/hr at boundary, 100 mR/hr inside RA",
            "Watchdog / standby person assigned during exposure",
            "Source-position verification — collimator + film / DR setup",
            "Cyclic monitoring of dose rate during shoot",
            "Source return-to-position confirmed via dual indicator (mechanical + survey)",
            "Daily survey log completed and signed",
            "Source storage — locked, ventilated, shielded enclosure post-shift",
        ]),
        ("ELECTRICAL SAFETY", [
            "GFCI / RCD on all 110V / 230V site power",
            "Cables inspected — no damage, kinks, exposed conductors",
            "Equipment grounded + bonded where applicable",
            "Generators positioned with adequate ventilation",
            "Magnetic particle yoke + prods — insulation intact",
            "Eddy current bench / portable — earthing verified",
            "Battery banks (for portable UT scopes) — vented, secured",
            "Hot work permit (welding / grinding) if applicable",
        ]),
        ("CHEMICAL SAFETY (PT / MT / Coatings)", [
            "SDS available for all chemicals on-site",
            "Solvent removable penetrant — flash point + ventilation OK",
            "Aerosol cans — date checked, no overheating",
            "Eyewash + emergency shower within 10 sec of chemical use",
            "Skin protection — nitrile / neoprene gloves per SDS",
            "Respiratory protection where required — fit-test current",
            "Waste container labeled — hazardous waste segregation",
            "Spill response materials staged near chemical use",
        ]),
        ("CONFINED SPACE / HEIGHT / OFFSHORE", [
            "Confined space entry permit issued — atmosphere tested",
            "Continuous gas monitoring (O₂, H₂S, LEL, CO) inside CS",
            "Rescue plan + retrieval gear on-site",
            "Attendant outside CS — log entries / exits",
            "Working at heights — fall protection (harness, lanyard, anchor)",
            "Scaffold tagged and inspected for current shift",
            "Offshore mobilization — BOSIET / HUET / FOET current",
            "MEDIC / sea-survival currency verified",
            "Lifejacket + EPIRB / PLB for offshore work",
        ]),
    ]

    row = 5
    for stitle, items in sections:
        row += 1
        xlsx_section(ws, row, stitle, 6, color=BRAND, dark=True)
        row += 1
        xlsx_header(ws, row, ["#", "Item", "Pass / Fail / N/A", "Initial", "Date", "Comments / Action"])
        for i, item in enumerate(items, 1):
            row += 1
            xlsx_data_row(ws, row, [i, item, "", "", "", ""])

    row += 2
    xlsx_section(ws, row, "Crew Sign-off", 6, color=BRAND_LIGHT)
    row += 1
    xlsx_header(ws, row, ["Name", "Role", "Cert No.", "Signature", "Date", "Comments"])
    for _ in range(8):
        row += 1
        xlsx_data_row(ws, row, ["", "", "", "", "", ""])

    row += 2
    xlsx_brand_footer(ws, row, 6)
    wb.save(f"{OUT}/ndt-safety-checklist.xlsx")

# ============================================================
# 4) Training Requirements Matrix (XLSX)
# ============================================================
def make_training_matrix():
    wb = Workbook()
    ws = wb.active
    ws.title = "Training Matrix"
    xlsx_set_widths(ws, [22, 14, 18, 18, 18, 22, 22])

    xlsx_title_row(ws, 1, "NDT Training Requirements Matrix — ASNT SNT-TC-1A / ISO 9712 / PCN", 7)
    xlsx_header(ws, 3, ["Method", "Level", "ASNT Training (hrs)", "ASNT Experience (hrs)", "ISO 9712 Training (hrs)", "ISO 9712 Min Experience (months)", "PCN Equivalent Requirements"])

    data = [
        ("Ultrasonic Testing (UT)", "Level I", 40, 210, 40, 3, "PCN UT Level 1 — 40 hr training + 3 mo OJT"),
        ("Ultrasonic Testing (UT)", "Level II", 40, 630, 80, 9, "PCN UT Level 2 — 80 hr training + 9 mo OJT"),
        ("Ultrasonic Testing (UT)", "Level III", "Per CP-189", "≥ 4 yr", "Per ISO 9712", "≥ 4 yr", "PCN Level 3 — ASNT/ISO eq."),
        ("Radiographic Testing (RT)", "Level I", 40, 210, 40, 3, "PCN RT Level 1 — 40 hr + 3 mo OJT"),
        ("Radiographic Testing (RT)", "Level II", 40, 630, 80, 9, "PCN RT Level 2 — 80 hr + 9 mo OJT"),
        ("Magnetic Particle (MT)", "Level I", 12, 130, 16, 1, "PCN MT Level 1"),
        ("Magnetic Particle (MT)", "Level II", 24, 390, 16, 3, "PCN MT Level 2"),
        ("Liquid Penetrant (PT)", "Level I", 8, 130, 16, 1, "PCN PT Level 1"),
        ("Liquid Penetrant (PT)", "Level II", 16, 390, 16, 3, "PCN PT Level 2"),
        ("Eddy Current (ET)", "Level I", 40, 210, 40, 3, "PCN ET Level 1"),
        ("Eddy Current (ET)", "Level II", 40, 630, 80, 9, "PCN ET Level 2"),
        ("Visual Testing (VT)", "Level I", 8, 130, 16, 1, "PCN VT Level 1"),
        ("Visual Testing (VT)", "Level II", 16, 390, 16, 3, "PCN VT Level 2"),
        ("Phased Array UT (PAUT)", "Level II", 80, 630, 80, 9, "PCN PAUT 80 hr + 9 mo OJT"),
        ("Time of Flight Diffraction (TOFD)", "Level II", 40, 320, 40, 6, "PCN TOFD 40 hr + 6 mo OJT"),
        ("Long Range UT (LRUT/GWT)", "Level II", 40, 320, 40, 6, "Per scheme certification"),
        ("AUT (Automated UT)", "Level II", 40, 320, 40, 6, "PCN AUT competency"),
        ("Computed Radiography (CR)", "Level II", 16, 210, 16, 3, "Add-on to RT Level II"),
        ("Digital Radiography (DR)", "Level II", 16, 210, 16, 3, "Add-on to RT Level II"),
    ]
    for i, row in enumerate(data, 4):
        xlsx_data_row(ws, i, row)
    for c in range(1, 8):
        ws.cell(row=4, column=c).font = BOLD if c == 1 else Font(name="Calibri", size=10)

    # Notes section
    n = 4 + len(data) + 2
    xlsx_section(ws, n, "Notes", 7, color=BRAND_LIGHT)
    notes = [
        "ASNT SNT-TC-1A 2024 edition values shown; CP-189-2020 may apply for some employers.",
        "ISO 9712:2021 training-hour minima shown; PCN GEN Issue 12 may impose additional requirements.",
        "Vision acuity (Jaeger J1/J2) + color discrimination annually for all certified personnel.",
        "Re-certification: ASNT 5-yr employer recert / ACCP 5-yr ASNT recert / ISO 9712 5-yr / PCN 5-yr.",
        "Sector / industry add-ons may require additional training (NAS-410, Boeing D-590, Saudi Aramco SAEP-1142, ADNOC ACS).",
        "Levels III: 4-year minimum experience including 1 year in the specific method.",
    ]
    for i, note in enumerate(notes, 1):
        ws.cell(row=n + i, column=1, value=f"  • {note}").alignment = LEFT
        ws.merge_cells(start_row=n + i, end_row=n + i, start_column=1, end_column=7)

    xlsx_brand_footer(ws, n + len(notes) + 2, 7)
    wb.save(f"{OUT}/training-requirements-matrix.xlsx")

# ============================================================
# DOCX helpers
# ============================================================
def docx_setup(doc, title):
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    style = doc.styles['Normal']
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    # Header
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    # Sub-line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Atlantis NDT — Free Editable Template")
    r.italic = True
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    r.font.size = Pt(10)
    doc.add_paragraph()

def docx_section(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def docx_label(doc, label, value=""):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    p.add_run(value)

def docx_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell._tc.get_or_add_tcPr().append(
            __import__('docx.oxml.ns', fromlist=['qn']).qn  # placeholder
        )  # will fail silently - styling minimal
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            table.rows[ri].cells[ci].text = str(val)

def docx_footer(doc):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("© Atlantis NDT — Free editable template. Customize for your method, code, and customer requirements. "
                  "Demo: info@atlantisndt.com  •  atlantisndt.com")
    r.italic = True
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    r.font.size = Pt(9)

# ============================================================
# 5) NDT Procedure Template (DOCX)
# ============================================================
def make_procedure_template():
    doc = Document()
    docx_setup(doc, "NDT Procedure Template — SNT-TC-1A / ISO 9712 Compliant")

    docx_section(doc, "1. Procedure Identification")
    for k, v in [
        ("Procedure Number", "[XXX-NDT-001]"),
        ("Revision", "[Rev 0]"),
        ("Effective Date", "[YYYY-MM-DD]"),
        ("Method", "[UT / RT / MT / PT / ET / VT / PAUT / TOFD]"),
        ("Application", "[e.g., Weld Inspection, Thickness Measurement, Tube Bundle Inspection]"),
        ("Material", "[CS / SS / Duplex / Cu / Al / Composite]"),
        ("Thickness Range", "[e.g., 6 mm – 100 mm]"),
        ("Geometry", "[e.g., Butt Weld, Pipe, Plate, Tube, Tank Shell]"),
        ("Reference Codes", "[ASME Section V Article 4 / API 510 / EN ISO 17640 / AWS D1.1]"),
        ("Customer / Project", "[Owner-Operator / EPC / Project No.]"),
        ("Prepared by", "[Name, ASNT Level III Cert No.]"),
        ("Reviewed by", "[Name, Position]"),
        ("Approved by", "[Name, Position]"),
    ]:
        docx_label(doc, k, v)

    docx_section(doc, "2. Scope")
    doc.add_paragraph(
        "This procedure defines the requirements for [Method] examination of [Application] in accordance with "
        "[Reference Code(s)] for [Customer]. The procedure applies to [Material] components in the thickness "
        "range of [Thickness Range] and geometry [Geometry]. Examination shall be performed by personnel "
        "qualified per [ASNT SNT-TC-1A / ISO 9712 / PCN] to the appropriate Level for the activity being "
        "performed."
    )

    docx_section(doc, "3. Reference Documents")
    refs = [
        "ASNT SNT-TC-1A (latest edition) — Recommended Practice for Personnel Qualification and Certification",
        "ASNT CP-189 (latest edition) — Standard for Qualification and Certification (if applicable)",
        "ISO 9712:2021 — NDT Personnel Qualification and Certification",
        "ASME Section V (latest edition) — NDE",
        "ASME Section VIII Division 1 — Pressure Vessels (acceptance criteria)",
        "ASME B31.3 — Process Piping (acceptance criteria)",
        "API 510 / 570 / 653 / 580 / 581 — In-service inspection codes",
        "[Customer Written Practice / Saudi Aramco SAEP / ADNOC ACS / Shell DEP — as applicable]",
        "Equipment manufacturer manual — [insert]",
        "[Calibration / reference block standard]",
    ]
    for r in refs:
        doc.add_paragraph(r, style="List Bullet")

    docx_section(doc, "4. Personnel Qualification")
    doc.add_paragraph(
        "Personnel performing the examination shall be qualified and certified to the appropriate Level per "
        "the employer's Written Practice [WPN-NDT-001] or per [ISO 9712 / PCN] as applicable. Level I "
        "personnel shall work under the supervision of qualified Level II or Level III. Level II personnel "
        "perform the examination, evaluate results, and prepare reports. Level III personnel are responsible "
        "for procedure approval, technique development, and final evaluation of indications."
    )
    docx_label(doc, "Minimum Level for Examination", "[Level II]")
    docx_label(doc, "Procedure Approval Level", "[Level III]")
    docx_label(doc, "Vision Acuity", "Annual Jaeger J1/J2 + Ishihara color")

    docx_section(doc, "5. Equipment & Materials")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Item", "Specification", "Make / Model / Standard", "Calibration / Verification"]):
        hdr[i].text = h
    for row_data in [
        ("Instrument", "[e.g., Digital UT flaw detector]", "[Olympus EPOCH 650 / Sonatest D70]", "Manufacturer functional check + linearity per ASME V"),
        ("Probe / Source", "[e.g., 5 MHz, Ø10 mm angle beam 60° / 45°]", "[Olympus / Eddyfi / Sonatest]", "Per ASME V Article 4 setup"),
        ("Couplant", "[ASTM E114 acceptable; e.g., Sonotech Ultragel II]", "[Manufacturer]", "Compatible with material"),
        ("Reference Block", "[IIW V1 / V2 / ASME basic / customer-specific]", "[Steel block, traceable]", "Per ASME V Annex"),
        ("Calibration Block", "[Step-wedge / DAC block]", "[As applicable]", "Verify before each scan"),
        ("Other Equipment", "[Survey meter / film badge / penetrant kit]", "[As applicable]", "Current calibration"),
    ]:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val

    docx_section(doc, "6. Surface Preparation")
    doc.add_paragraph(
        "Inspection surfaces shall be free of scale, dirt, weld spatter, paint, and any condition that would "
        "interfere with the examination. Surface preparation requirements:"
    )
    for item in [
        "Visual inspection: degrease and clean per ASME V Article 9 §T-921",
        "Magnetic Particle: clean to bare metal per ASTM E1444 §6",
        "Liquid Penetrant: degrease, dry; coating-free per ASME V Article 6 §T-621",
        "Ultrasonic: smooth surface, Ra ≤ 6.3 µm; remove scale and loose material",
        "Radiographic: surface need only allow source/film/IP placement; weld-cap not required to be flush",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    docx_section(doc, "7. Examination Technique")
    doc.add_paragraph(
        "Examination shall be conducted as follows. The technique includes calibration / system verification, "
        "scan plan with coverage % calculation, indication detection, indication characterization, sizing, "
        "and recording."
    )
    docx_section(doc, "7.1 Calibration / System Verification", level=2)
    doc.add_paragraph("[Per ASME V Article 4 §T-460 / Annex VIII for UT — describe the calibration sequence and DAC / DGS / TCG setup, sensitivity, and verification frequency.]")
    docx_section(doc, "7.2 Scan Plan & Coverage", level=2)
    doc.add_paragraph("[Describe scan increments, overlap %, scan speed limits, probe coupling check intervals.]")
    docx_section(doc, "7.3 Indication Detection & Characterization", level=2)
    doc.add_paragraph("[Define recording threshold (e.g., DAC -6 dB), characterization criteria (planar vs volumetric), and sizing methodology (6 dB drop / 20 dB drop / amplitude / TOFD tip-diffraction).]")

    docx_section(doc, "8. Acceptance Criteria")
    doc.add_paragraph(
        "Indications shall be evaluated against the acceptance criteria of the applicable construction code "
        "or customer specification. Where multiple criteria apply, the most stringent shall govern."
    )
    docx_label(doc, "Primary Code", "[ASME Section VIII Division 1 / B31.3 / API 510 / AWS D1.1]")
    docx_label(doc, "Customer Specification", "[if applicable]")
    docx_label(doc, "Stricter Limit Applied", "[Reference + clause]")

    docx_section(doc, "9. Reporting")
    doc.add_paragraph(
        "Inspection results shall be documented on a format that includes the following minimum elements:"
    )
    for item in [
        "Procedure number, revision, and date",
        "Component identification (asset, weld map ID, drawing reference)",
        "Examination method, technique, and acceptance code reference",
        "Equipment used (instrument, probe, calibration block, source) with serial numbers",
        "Inspector name and qualification (Level + ASNT/ISO/PCN cert no.)",
        "Calibration record at start and end of examination",
        "All indications with location, size, characterization, and acceptance result",
        "Accept / reject decision per acceptance criteria",
        "Disposition for rejects (repair / rework / waiver request)",
        "Photographs / images / data files (where applicable)",
    ]:
        doc.add_paragraph(item, style="List Number")

    docx_section(doc, "10. Records & Retention")
    doc.add_paragraph(
        "All examination records, calibration certificates, personnel certification copies, and technique "
        "sheets shall be retained for [7 years] post-project closeout or as specified by [Customer / "
        "Regulatory authority]. Records shall be stored in a controlled environment with retrieval traceability "
        "per ISO 9001 §7.5 or ISO/IEC 17025 §8.4."
    )

    docx_section(doc, "11. Revision History")
    table = doc.add_table(rows=2, cols=4)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Rev", "Date", "Description", "Approved by"]):
        table.rows[0].cells[i].text = h
    for i, v in enumerate(["0", "[YYYY-MM-DD]", "Initial issue", "[Level III Name]"]):
        table.rows[1].cells[i].text = v

    docx_footer(doc)
    doc.save(f"{OUT}/ndt-procedure-template.docx")

# ============================================================
# 6) ASNT Level III Study Guide (DOCX)
# ============================================================
def make_level_iii_study_guide():
    doc = Document()
    docx_setup(doc, "ASNT NDT Level III — Exam Study Guide")

    p = doc.add_paragraph()
    r = p.add_run("Use this study guide for ASNT NDT Level III Basic and Method exam preparation. Covers exam format, topic coverage by method, key reference standards, recommended study timeline, and test-day strategies.")
    r.italic = True

    docx_section(doc, "1. ASNT Level III Exam Structure")
    doc.add_paragraph(
        "The ASNT Level III qualification consists of: (1) Basic Exam — 135 questions covering NDT fundamentals, "
        "materials science, manufacturing processes, and quality concepts; (2) Method Exam — 66 questions on the "
        "specific NDT method (UT / RT / MT / PT / ET / VT / LT / NR / AE); (3) Specific Examination — employer-"
        "administered exam covering codes, standards, and procedures applicable to the employer's products."
    )
    docx_section(doc, "1.1 Time Allowed", level=2)
    docx_label(doc, "Basic Exam", "2.5 hours, closed-book")
    docx_label(doc, "Method Exam", "2 hours, closed-book")
    docx_label(doc, "Specific Exam", "Variable, employer-defined")
    docx_section(doc, "1.2 Pass Threshold", level=2)
    docx_label(doc, "Each Exam", "≥ 70% composite (some methods 80%)")

    docx_section(doc, "2. Basic Exam — Topic Coverage")
    for item in [
        "Materials science: ferrous and non-ferrous metals, polymers, composites, ceramics",
        "Manufacturing processes: casting, forging, welding, rolling, extrusion, machining",
        "Discontinuities: inherent (casting porosity, hot tears), processing (lamination, hydrogen flakes), service-induced (fatigue, corrosion, SCC)",
        "Materials properties: tensile, hardness, impact, fracture toughness",
        "NDT method overview — capabilities, limitations, application areas of each method",
        "Quality concepts: SPC, AQL, sampling plans, gauge R&R, ISO 9001 / AS9100 fundamentals",
        "Codes and standards overview — ASME V, API ICP, EN ISO, AWS D1.1, MIL-STD",
        "Personnel qualification — SNT-TC-1A, CP-189, ACCP, ISO 9712, NAS-410",
        "Safety — radiation, electrical, chemical, confined space",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    docx_section(doc, "3. Method Exam — Topic Coverage (UT example)")
    for item in [
        "Wave propagation: longitudinal, shear, surface, Lamb modes, mode conversion",
        "Probes: piezoelectric, phased array, EMAT — frequency, element size, beam profile",
        "Calibration: DAC, DGS, TCG, time-corrected gain, reference block selection (IIW V1/V2, ASME basic)",
        "Scanning techniques: contact, immersion, automated, phased array, TOFD",
        "Discontinuity characterization: planar vs volumetric, tip-diffraction sizing, 6/20 dB drop",
        "Specific applications: weld inspection, thickness measurement, tube bundle inspection, casting / forging",
        "Code requirements: ASME V Article 4 (UT for welds), Article 5 (other), Annex VIII (PAUT)",
        "Equipment performance verification: ASME V Annex II, linearity, dead zone, near field",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    docx_section(doc, "4. Recommended Study Timeline")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Weeks", "Focus", "Activities"]):
        table.rows[0].cells[i].text = h
    timeline = [
        ("Week 1-2", "Foundations", "Read ASNT Level III Study Guide (Method); review codes; build summary notes."),
        ("Week 3-4", "Materials & Processes", "Materials science chapters; manufacturing processes; defect types."),
        ("Week 5-6", "Method Deep-Dive", "Method-specific chapter; calibration techniques; code applications."),
        ("Week 7", "Practice Questions", "ASNT Level III Question Books; 200+ practice questions."),
        ("Week 8", "Mock Exam + Weak Areas", "2 full mock exams; identify gaps; targeted study."),
        ("Week 9", "Code Drill", "ASME V / API 510/570/653 / customer specs — code-tab strategy."),
        ("Week 10", "Final Review + Exam", "Light review; rest; exam day."),
    ]
    for row_data in timeline:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val

    docx_section(doc, "5. Reference Materials")
    for item in [
        "ASNT Level III Study Guide — Basic + Method (per method)",
        "ASNT Question Books — Level III (per method)",
        "ASME Section V (current edition) — NDE",
        "ASME Section IX — Welding & Brazing Qualifications",
        "API 510 / 570 / 653 / 580 / 581 Body of Knowledge",
        "AWS D1.1 — Structural Welding Code",
        "Atlantis NDT Level III prep course materials (Houston, Dubai, Hyderabad, Online)",
        "ASNT Handbook (current edition) — for the specific method",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    docx_section(doc, "6. Test-Day Strategy")
    for item in [
        "Sleep 7-8 hours night before. No new material that day.",
        "Eat a light protein-rich breakfast.",
        "Arrive 30 min early. Bring two photo IDs.",
        "Read the question twice before looking at answers. Eliminate clearly-wrong answers first.",
        "Skip hard questions; flag and return after completing easier ones.",
        "Time management: ~1 min per question for Basic. Method exam more time per question.",
        "Don't leave any question blank. Educated guess > blank.",
        "Review flagged questions if time permits. Trust first instinct unless you find new info.",
        "Stay calm. The exam is not designed to trick you — read the question precisely.",
    ]:
        doc.add_paragraph(item, style="List Number")

    docx_section(doc, "7. After the Exam")
    doc.add_paragraph(
        "ASNT typically notifies results within 4-6 weeks. Passing both Basic and Method qualifies you as "
        "ASNT NDT Level III for that method. Apply for the employer-specific written practice and complete "
        "any Specific Examination required by your employer. Maintain certification with annual review per "
        "SNT-TC-1A or 5-yearly recertification per ACCP / ISO 9712 / PCN."
    )
    doc.add_paragraph(
        "If you fail, ASNT permits retest after 30 days. Most candidates who fail do so on Basic — invest "
        "additional time in materials science, manufacturing processes, and quality concepts. Method failures "
        "often relate to underestimating the depth of code knowledge required."
    )

    docx_section(doc, "8. Atlantis NDT Level III Prep Course")
    doc.add_paragraph(
        "Atlantis NDT delivers ASNT Level III preparation in Houston (HQ), Dubai, Hyderabad, and online. "
        "Course duration: 80-120 hours per method depending on candidate background. 95% first-attempt pass "
        "rate. ASNT Level III instructors with multi-method certification and active inspection experience. "
        "Mock exams, code-tab strategy, and 1-on-1 mentoring included. Contact info@atlantisndt.com to enroll."
    )

    docx_footer(doc)
    doc.save(f"{OUT}/asnt-level-iii-study-guide.docx")

# ============================================================
# 7) API 510 Pressure Vessel Inspection Report (XLSX) — multi-sheet
# ============================================================
def make_api510_report():
    wb = Workbook()

    # Cover
    cv = wb.active
    cv.title = "Cover"
    xlsx_set_widths(cv, [4, 32, 32, 32])
    xlsx_title_row(cv, 1, "API 510 Pressure Vessel Inspection Record", 4)
    info = [
        ("Vessel Tag No.", "Vessel Description"),
        ("Site / Unit", "Inspection Date"),
        ("Design Code", "Year of Construction"),
        ("National Board No.", "Manufacturer / Serial No."),
        ("Design Pressure (psi / kPa)", "Design Temperature (deg F / C)"),
        ("MAWP (psi / kPa)", "MDMT (deg F / C)"),
        ("Service / Lethal (Y/N)", "Specific Gravity / Fluid"),
        ("Shell Material / Spec.", "Head Material / Spec."),
        ("Internal Diameter (in / mm)", "Shell Thickness (nominal)"),
        ("Internal Lining (Y/N + type)", "External Insulation (Y/N + type)"),
        ("PWHT (Y/N)", "Last Internal Inspection Date"),
        ("Next Internal Due", "Last External Inspection Date"),
        ("Next External Due", "RBI Interval (months)"),
        ("Inspector Name", "API 510 Cert No."),
        ("Inspection Type", "Inspection Result (Pass / Conditional / Fail)"),
    ]
    r = 3
    for left, right in info:
        cv.cell(row=r, column=1, value=left).font = BOLD
        cv.cell(row=r, column=2, value="").border = BORDER
        cv.cell(row=r, column=3, value=right).font = BOLD
        cv.cell(row=r, column=4, value="").border = BORDER
        r += 1
    r += 2
    xlsx_brand_footer(cv, r, 4)

    # Shell UT Thickness
    sh = wb.create_sheet("Shell UT Thickness")
    xlsx_set_widths(sh, [10, 14, 14, 14, 14, 14, 14, 14, 14, 18])
    xlsx_title_row(sh, 1, "API 510 Shell / Head UT Thickness — CML Readings", 10)
    xlsx_header(sh, 3, ["CML ID", "Location", "Nominal t (in)", "Reading 1", "Reading 2", "Reading 3", "Reading 4", "Min t (in)", "Avg t (in)", "Corrosion Rate (mpy)"])
    for i in range(1, 21):
        sh.cell(row=3 + i, column=1, value=f"CML-{i:03d}")
        for c in range(2, 11):
            sh.cell(row=3 + i, column=c).border = BORDER
            sh.cell(row=3 + i, column=c).alignment = CENTER if c > 2 else LEFT
            sh.cell(row=3 + i, column=c).font = Font(name="Calibri", size=10)
        sh.cell(row=3 + i, column=1).border = BORDER
        sh.cell(row=3 + i, column=1).font = Font(name="Calibri", size=10, bold=True)
        sh.row_dimensions[3 + i].height = 20

    # Nozzle Inspection
    nz = wb.create_sheet("Nozzle Inspection")
    xlsx_set_widths(nz, [10, 18, 14, 14, 14, 14, 14, 22])
    xlsx_title_row(nz, 1, "Nozzle / Manway Inspection", 8)
    xlsx_header(nz, 3, ["Nozzle ID", "Service / Description", "Size (NPS)", "Reinf. Pad (Y/N)", "Min t (in)", "Required t", "Result (P/F)", "Comments"])
    nozzles = ["N1 (Inlet)", "N2 (Outlet)", "N3 (Vent)", "N4 (Drain)", "N5 (PSV)", "N6 (Level Gauge)", "N7 (Instrument)", "N8 (Spare)", "M1 (Manway Top)", "M2 (Manway Bottom)"]
    for i, name in enumerate(nozzles, 1):
        xlsx_data_row(nz, 3 + i, [name, "", "", "", "", "", "", ""])

    # Weld Inspection
    wl = wb.create_sheet("Weld Inspection")
    xlsx_set_widths(wl, [10, 22, 14, 14, 16, 14, 14, 22])
    xlsx_title_row(wl, 1, "Long Seam / Circ Seam Weld Inspection", 8)
    xlsx_header(wl, 3, ["Weld ID", "Location / Joint", "Method", "Length (in)", "Coverage %", "Indications", "Accept/Reject", "Reference / Comments"])
    weld_types = ["LS-1", "LS-2", "LS-3", "CS-1", "CS-2", "CS-3", "CS-4", "Head A-Shell", "Head B-Shell", "Nozzle N1 weld"]
    for i, w in enumerate(weld_types, 1):
        xlsx_data_row(wl, 3 + i, [w, "", "", "", "", "", "", ""])

    # FFS Screening (API 579)
    ff = wb.create_sheet("FFS Screening")
    xlsx_set_widths(ff, [4, 30, 16, 16, 16, 14, 22])
    xlsx_title_row(ff, 1, "Fitness-for-Service Screening — API 579-1 / ASME FFS-1", 7)
    xlsx_header(ff, 3, ["#", "Damage Mechanism", "Found (Y/N)", "Level 1 Pass?", "Level 2 Required?", "Action", "Comments"])
    mechs = [
        "General / Uniform metal loss",
        "Local thin area (LTA)",
        "Pitting damage",
        "Blisters / HIC / SOHIC",
        "Weld misalignment / shell distortion",
        "Crack-like flaw",
        "Creep damage",
        "Fire damage",
        "Dent / gouge",
        "Laminations",
        "Pitting (atmospheric / CUI)",
        "Brittle fracture concern (MDMT)",
    ]
    for i, m in enumerate(mechs, 1):
        xlsx_data_row(ff, 3 + i, [i, m, "", "", "", "", ""])

    # Recommendations
    rc = wb.create_sheet("Recommendations")
    xlsx_set_widths(rc, [4, 14, 50, 14, 14, 20])
    xlsx_title_row(rc, 1, "Recommendations & Follow-up Actions", 6)
    xlsx_header(rc, 3, ["#", "Priority", "Description", "API 510 Ref.", "Due Date", "Responsible"])
    for i in range(1, 11):
        xlsx_data_row(rc, 3 + i, [i, "", "", "", "", ""])

    # Sign-off
    so = wb.create_sheet("Sign-off")
    xlsx_set_widths(so, [4, 28, 22, 18, 14])
    xlsx_title_row(so, 1, "Sign-off", 5)
    xlsx_header(so, 3, ["#", "Role", "Name", "API 510 Cert No.", "Date"])
    for i, role in enumerate(["API 510 Inspector", "Owner/User Engineer", "Independent Reviewer", "Jurisdiction (if applicable)"], 1):
        xlsx_data_row(so, 3 + i, [i, role, "", "", ""])
    xlsx_brand_footer(so, 10, 5)
    wb.save(f"{OUT}/api-510-inspection-report.xlsx")

# ============================================================
# 8) API 570 Piping Inspection Record (XLSX)
# ============================================================
def make_api570_record():
    wb = Workbook()

    cv = wb.active
    cv.title = "Cover"
    xlsx_set_widths(cv, [4, 32, 32, 32])
    xlsx_title_row(cv, 1, "API 570 Piping Circuit Inspection Record", 4)
    info = [
        ("Piping Circuit ID", "P&ID Reference"),
        ("Site / Unit", "Inspection Date"),
        ("Service / Fluid", "Class (1/2/3)"),
        ("Design Pressure (psi)", "Design Temperature (deg F)"),
        ("Operating Pressure (psi)", "Operating Temperature (deg F)"),
        ("Pipe Material / Spec.", "Pipe Size / Schedule"),
        ("Insulation (Y/N + type)", "PWHT (Y/N)"),
        ("Last Thickness Survey", "Next Thickness Survey Due"),
        ("Last External Visual", "Next External Visual Due"),
        ("RBI Interval (months)", "Inspection Effectiveness (A/B/C/D)"),
        ("Inspector Name", "API 570 Cert No."),
        ("Inspection Type", "Inspection Result"),
    ]
    r = 3
    for left, right in info:
        cv.cell(row=r, column=1, value=left).font = BOLD
        cv.cell(row=r, column=2, value="").border = BORDER
        cv.cell(row=r, column=3, value=right).font = BOLD
        cv.cell(row=r, column=4, value="").border = BORDER
        r += 1
    r += 2
    xlsx_brand_footer(cv, r, 4)

    # CML Thickness Readings
    cm = wb.create_sheet("CML Readings")
    xlsx_set_widths(cm, [10, 18, 14, 14, 14, 14, 14, 14, 14, 16])
    xlsx_title_row(cm, 1, "Condition Monitoring Locations (CML) — Thickness Survey", 10)
    xlsx_header(cm, 3, ["CML ID", "Location / ISO Sketch", "Nominal t (in)", "Previous t (in)", "Current t (in)", "Min t Required", "ST CR (mpy)", "LT CR (mpy)", "Remaining Life (yr)", "Status (P/F)"])
    for i in range(1, 31):
        cm.cell(row=3 + i, column=1, value=f"CML-{i:03d}")
        cm.cell(row=3 + i, column=1).border = BORDER
        cm.cell(row=3 + i, column=1).font = Font(name="Calibri", size=10, bold=True)
        for c in range(2, 11):
            cm.cell(row=3 + i, column=c).border = BORDER
            cm.cell(row=3 + i, column=c).alignment = CENTER if c > 2 else LEFT
            cm.cell(row=3 + i, column=c).font = Font(name="Calibri", size=10)
        cm.row_dimensions[3 + i].height = 20

    # Corrosion Rate / Remaining Life summary
    cr = wb.create_sheet("CR & Remaining Life")
    xlsx_set_widths(cr, [4, 30, 16, 16, 16, 16, 16])
    xlsx_title_row(cr, 1, "Corrosion Rate Summary & Remaining Life", 7)
    xlsx_header(cr, 3, ["#", "Section / Component", "Nominal (in)", "Min Found (in)", "Min Req. (in)", "Worst CR (mpy)", "Remaining Life (yr)"])
    sections = [
        "Straight Run — Sec A",
        "Straight Run — Sec B",
        "Straight Run — Sec C",
        "Elbow E-1 (extrados)",
        "Elbow E-2 (intrados)",
        "Tee T-1 (header)",
        "Tee T-2 (branch)",
        "Reducer R-1",
        "Flange F-1 (raised face)",
        "Block Valve BV-1",
        "Dead leg DL-1",
        "Sample point SP-1",
    ]
    for i, s in enumerate(sections, 1):
        xlsx_data_row(cr, 3 + i, [i, s, "", "", "", "", ""])

    # RBI Worksheet (light)
    rb = wb.create_sheet("RBI Score")
    xlsx_set_widths(rb, [4, 28, 14, 14, 14, 14, 22])
    xlsx_title_row(rb, 1, "Risk Score — API 580 / 581 Simplified", 7)
    xlsx_header(rb, 3, ["#", "Damage Mechanism", "POF (1-5)", "COF Safety", "COF Env.", "COF Financial", "Risk Cell (e.g., 3C)"])
    dm = [
        "Internal corrosion (general)",
        "CUI (insulated)",
        "Erosion / erosion-corrosion",
        "Microbiologically induced corrosion (MIC)",
        "Chloride SCC",
        "Sulfide SCC / HIC",
        "Thermal fatigue",
        "Mechanical fatigue",
        "Brittle fracture",
        "Creep (high temp lines)",
    ]
    for i, m in enumerate(dm, 1):
        xlsx_data_row(rb, 3 + i, [i, m, "", "", "", "", ""])

    # Recommendations
    rc = wb.create_sheet("Recommendations")
    xlsx_set_widths(rc, [4, 14, 50, 14, 14, 20])
    xlsx_title_row(rc, 1, "Recommendations / Repairs / Re-inspection", 6)
    xlsx_header(rc, 3, ["#", "Priority", "Description", "API 570 Ref.", "Due Date", "Responsible"])
    for i in range(1, 11):
        xlsx_data_row(rc, 3 + i, [i, "", "", "", "", ""])

    # Sign-off
    so = wb.create_sheet("Sign-off")
    xlsx_set_widths(so, [4, 28, 22, 18, 14])
    xlsx_title_row(so, 1, "Sign-off", 5)
    xlsx_header(so, 3, ["#", "Role", "Name", "API 570 Cert No.", "Date"])
    for i, role in enumerate(["API 570 Inspector", "Piping Engineer", "RBI Specialist", "Owner/User"], 1):
        xlsx_data_row(so, 3 + i, [i, role, "", "", ""])
    xlsx_brand_footer(so, 10, 5)

    wb.save(f"{OUT}/api-570-piping-inspection-record.xlsx")

# ============================================================
# 9) Daily Progress Report (DPR) (XLSX)
# ============================================================
def make_dpr():
    wb = Workbook()
    ws = wb.active
    ws.title = "DPR"
    xlsx_set_widths(ws, [4, 36, 14, 14, 14, 14, 22])

    xlsx_title_row(ws, 1, "Daily Progress Report — NDT Inspection Project", 7)
    ws.cell(row=2, column=1, value="Project:").font = BOLD
    ws.cell(row=2, column=4, value="Date:").font = BOLD
    ws.cell(row=3, column=1, value="Client:").font = BOLD
    ws.cell(row=3, column=4, value="Site / Unit:").font = BOLD
    ws.cell(row=4, column=1, value="DPR No.:").font = BOLD
    ws.cell(row=4, column=4, value="Shift (Day / Night):").font = BOLD
    ws.cell(row=5, column=1, value="Project Manager:").font = BOLD
    ws.cell(row=5, column=4, value="Crew Lead / QC:").font = BOLD
    ws.cell(row=6, column=1, value="Weather:").font = BOLD
    ws.cell(row=6, column=4, value="Ambient Temp.:").font = BOLD

    # Manpower
    row = 8
    xlsx_section(ws, row, "Manpower Deployed Today", 7, color=BRAND, dark=True)
    row += 1
    xlsx_header(ws, row, ["#", "Role", "Planned", "Actual", "Hours", "OT Hours", "Comments"])
    roles = ["Level III Supervisor", "Level II UT Tech", "Level II RT Tech", "Level II MT/PT Tech", "Level II PAUT Tech", "RT Assistant / Helper", "QC Inspector", "Safety Officer", "Driver / Crew Support"]
    for i, r in enumerate(roles, 1):
        row += 1
        xlsx_data_row(ws, row, [i, r, "", "", "", "", ""])

    # Inspection Counts by Method
    row += 2
    xlsx_section(ws, row, "Inspections Completed Today (by Method)", 7, color=BRAND, dark=True)
    row += 1
    xlsx_header(ws, row, ["#", "Method", "Planned (today)", "Completed", "Accepted", "Rejected", "Backlog / Hold"])
    methods = ["Ultrasonic Testing (UT)", "Phased Array UT (PAUT)", "TOFD", "Radiography (RT) — Gamma", "Radiography (RT) — X-Ray / DR", "Magnetic Particle (MT)", "Liquid Penetrant (PT)", "Visual (VT)", "Eddy Current (ET)", "Hardness Testing", "PMI", "Heat Treatment Witness"]
    for i, m in enumerate(methods, 1):
        row += 1
        xlsx_data_row(ws, row, [i, m, "", "", "", "", ""])

    # Holds & Photo Summary
    row += 2
    xlsx_section(ws, row, "Holds, Rejects & Blockers", 7, color=BRAND, dark=True)
    row += 1
    xlsx_header(ws, row, ["#", "Description (component / weld / area)", "Method", "Defect Type", "Action Required", "Owner", "ETA"])
    for i in range(1, 9):
        row += 1
        xlsx_data_row(ws, row, [i, "", "", "", "", "", ""])

    # HSE Incidents
    row += 2
    xlsx_section(ws, row, "HSE Summary", 7, color=BRAND, dark=True)
    row += 1
    xlsx_header(ws, row, ["#", "Incident / Observation", "Severity (LTI/MTC/FAC/NM)", "Investigation Started", "CAPA", "Owner", "Comments"])
    for i in range(1, 5):
        row += 1
        xlsx_data_row(ws, row, [i, "", "", "", "", "", ""])

    # Look-ahead
    row += 2
    xlsx_section(ws, row, "Look-Ahead (Next 24h)", 7, color=BRAND, dark=True)
    row += 1
    xlsx_header(ws, row, ["#", "Planned Activity", "Method", "Location / Asset", "Resources Required", "Hold Points", "Notes"])
    for i in range(1, 7):
        row += 1
        xlsx_data_row(ws, row, [i, "", "", "", "", "", ""])

    # Sign-off
    row += 2
    xlsx_section(ws, row, "Sign-off", 7, color=BRAND_LIGHT)
    row += 1
    xlsx_header(ws, row, ["Role", "Name", "Cert No. / ID", "Signature", "Date", "Time", "Comments"])
    for role in ["Crew Lead", "Project Engineer", "Client Representative", "HSE Officer"]:
        row += 1
        xlsx_data_row(ws, row, [role, "", "", "", "", "", ""])

    row += 2
    xlsx_brand_footer(ws, row, 7)
    wb.save(f"{OUT}/daily-progress-report-dpr.xlsx")

# ============================================================
# 10) PWHT Record (XLSX)
# ============================================================
def make_pwht_record():
    wb = Workbook()
    ws = wb.active
    ws.title = "PWHT Record"
    xlsx_set_widths(ws, [4, 30, 18, 14, 14, 14, 22])

    xlsx_title_row(ws, 1, "Post-Weld Heat Treatment Record — ASME B31.3 / Section VIII", 7)
    info = [
        ("PWHT Record No.", "Date of PWHT"),
        ("Project / Job No.", "Site / Location"),
        ("Component ID / Weld No.", "WPS / PQR Reference"),
        ("Material Spec. / P-No.", "Thickness (in / mm)"),
        ("Heat Treatment Method", "Heater Type (resistance / induction / furnace)"),
        ("Code of Construction", "Required PWHT Per (Code / Spec)"),
        ("Soak Temperature (deg F / C)", "Soak Tolerance (+/- deg F / C)"),
        ("Required Hold Time (hr)", "Heating Rate (max deg F/hr)"),
        ("Cooling Rate (max deg F/hr)", "Loading Temp. (max deg F / C)"),
        ("Chart Recorder Serial No.", "Last Calibration Date"),
        ("Number of Thermocouples", "Thermocouple Type (K/J/N)"),
        ("Insulation Used", "Soak Band Width (in / mm)"),
    ]
    r = 3
    for left, right in info:
        ws.cell(row=r, column=1, value=left).font = BOLD
        ws.merge_cells(start_row=r, end_row=r, start_column=2, end_column=4)
        ws.cell(row=r, column=2).border = BORDER
        ws.cell(row=r, column=5, value=right).font = BOLD
        ws.merge_cells(start_row=r, end_row=r, start_column=6, end_column=7)
        ws.cell(row=r, column=6).border = BORDER
        r += 1

    # TC layout
    r += 1
    xlsx_section(ws, r, "Thermocouple (TC) Layout & Placement", 7, color=BRAND, dark=True)
    r += 1
    xlsx_header(ws, r, ["TC No.", "Location on Component", "Side (Top / Bottom / Side)", "Distance from Weld (in)", "Attachment (CDW / strap)", "Channel", "Comments"])
    for i in range(1, 11):
        r += 1
        xlsx_data_row(ws, r, [f"TC-{i}", "", "", "", "", "", ""])

    # Cycle Data
    r += 2
    xlsx_section(ws, r, "Heat Treatment Cycle Data", 7, color=BRAND, dark=True)
    r += 1
    xlsx_header(ws, r, ["#", "Time / Phase", "Avg Temp (deg F)", "Max Temp (deg F)", "Min Temp (deg F)", "Rate (deg F/hr)", "Comments"])
    phases = ["Cold Start", "End of Heat Ramp", "Begin Soak", "Mid Soak", "End Soak", "End of Cool Ramp", "Below 600 deg F"]
    for i, p in enumerate(phases, 1):
        r += 1
        xlsx_data_row(ws, r, [i, p, "", "", "", "", ""])

    # Acceptance
    r += 2
    xlsx_section(ws, r, "Acceptance Check vs Code Requirements", 7, color=BRAND, dark=True)
    r += 1
    xlsx_header(ws, r, ["#", "Parameter", "Required", "Actual", "Pass / Fail", "Code Ref.", "Comments"])
    params = [
        "Soak Temperature within tolerance",
        "Hold Time at soak met",
        "Heating Rate within limit",
        "Cooling Rate within limit",
        "Max differential between TCs",
        "Loading / unloading temperature met",
        "Chart trace continuous (no gaps)",
        "Pre-PWHT NDT performed",
        "Post-PWHT NDT performed",
    ]
    for i, p in enumerate(params, 1):
        r += 1
        xlsx_data_row(ws, r, [i, p, "", "", "", "", ""])

    # Attachments
    r += 2
    xlsx_section(ws, r, "Attachments Checklist", 7, color=BRAND_LIGHT)
    r += 1
    xlsx_header(ws, r, ["#", "Document", "Attached (Y/N)", "Ref. / File No.", "Issued By", "Date", "Comments"])
    attachments = [
        "PWHT chart trace (printed + electronic)",
        "Calibration certificate — chart recorder",
        "Calibration certificate — thermocouples",
        "TC layout sketch",
        "WPS / PQR reference",
        "Material certificate / MTR",
        "Heat treatment crew certificates",
    ]
    for i, a in enumerate(attachments, 1):
        r += 1
        xlsx_data_row(ws, r, [i, a, "", "", "", "", ""])

    # Sign-off
    r += 2
    xlsx_section(ws, r, "Sign-off", 7, color=BRAND_LIGHT)
    r += 1
    xlsx_header(ws, r, ["Role", "Name", "Cert No.", "Signature", "Date", "Time", "Comments"])
    for role in ["PWHT Operator", "NDT Level II / III Witness", "QC Inspector", "Client / TPI"]:
        r += 1
        xlsx_data_row(ws, r, [role, "", "", "", "", "", ""])

    r += 2
    xlsx_brand_footer(ws, r, 7)
    wb.save(f"{OUT}/pwht-record.xlsx")

# ============================================================
# 11) RBI Worksheet (API 581) (XLSX)
# ============================================================
def make_rbi_worksheet():
    wb = Workbook()

    # Asset register
    ar = wb.active
    ar.title = "Asset Register"
    xlsx_set_widths(ar, [4, 18, 26, 14, 16, 16, 16, 18])
    xlsx_title_row(ar, 1, "Risk-Based Inspection (RBI) Worksheet — API 581", 8)
    xlsx_header(ar, 3, ["#", "Asset Tag", "Description / Service", "Equip. Type", "Design Code", "Year Installed", "Last Inspection", "Next Due"])
    for i in range(1, 16):
        xlsx_data_row(ar, 3 + i, [i, "", "", "", "", "", "", ""])

    # Damage Mechanism Assessment
    dm = wb.create_sheet("Damage Mechanisms")
    xlsx_set_widths(dm, [4, 16, 26, 16, 16, 16, 16, 16, 16])
    xlsx_title_row(dm, 1, "Damage Mechanism Assessment", 9)
    xlsx_header(dm, 3, ["#", "Asset Tag", "Damage Mechanism", "Susceptibility (H/M/L/N)", "Damage Factor (Df)", "Confidence", "Insp. Effectiveness (A-D)", "Df Updated", "Comments"])
    mechs = [
        "General / uniform corrosion",
        "Localized corrosion",
        "Pitting",
        "Erosion / erosion-corrosion",
        "MIC",
        "Cl-SCC (austenitic SS)",
        "PolySCC / Sulfide SCC",
        "Sulfide stress cracking (HIC/SOHIC)",
        "Carbonate cracking (CC)",
        "Caustic cracking",
        "Amine cracking",
        "Brittle fracture",
        "Creep / stress rupture",
        "Thermal fatigue",
        "Mechanical fatigue",
        "Hydrogen blistering",
        "Refractory degradation",
        "External CUI",
        "Atmospheric corrosion",
        "Other (specify)",
    ]
    for i, m in enumerate(mechs, 1):
        xlsx_data_row(dm, 3 + i, [i, "", m, "", "", "", "", "", ""])

    # POF / COF / Risk
    rk = wb.create_sheet("POF-COF-Risk")
    xlsx_set_widths(rk, [4, 18, 14, 14, 14, 14, 14, 14, 14, 16])
    xlsx_title_row(rk, 1, "POF / COF / Risk Calculation", 10)
    xlsx_header(rk, 3, ["#", "Asset Tag", "Generic Freq.", "Total Df", "POF (cat 1-5)", "COF Safety", "COF Env.", "COF Financial", "Composite COF (A-E)", "Risk Cell"])
    for i in range(1, 16):
        xlsx_data_row(rk, 3 + i, [i, "", "", "", "", "", "", "", "", ""])

    # Risk Matrix
    mt = wb.create_sheet("Risk Matrix")
    xlsx_set_widths(mt, [10, 14, 14, 14, 14, 14])
    xlsx_title_row(mt, 1, "Risk Matrix — API 581 5x5", 6)
    xlsx_header(mt, 3, ["POF \\ COF", "A (Low)", "B", "C", "D", "E (High)"])
    matrix = [
        ("5 (High)", "Medium", "High", "High", "Very High", "Very High"),
        ("4", "Medium", "Medium", "High", "High", "Very High"),
        ("3", "Low", "Medium", "Medium", "High", "High"),
        ("2", "Low", "Low", "Medium", "Medium", "High"),
        ("1 (Low)", "Low", "Low", "Low", "Medium", "Medium"),
    ]
    for i, row in enumerate(matrix, 1):
        xlsx_data_row(mt, 3 + i, list(row))

    # Inspection Plan
    ip = wb.create_sheet("Inspection Plan")
    xlsx_set_widths(ip, [4, 18, 22, 14, 14, 14, 14, 22])
    xlsx_title_row(ip, 1, "Risk-Mitigation Inspection Plan", 8)
    xlsx_header(ip, 3, ["#", "Asset Tag", "Damage Mechanism", "Recommended NDT", "Effectiveness Target", "Coverage %", "Due Date", "Comments"])
    for i in range(1, 16):
        xlsx_data_row(ip, 3 + i, [i, "", "", "", "", "", "", ""])

    # Sign-off
    so = wb.create_sheet("Sign-off")
    xlsx_set_widths(so, [4, 28, 22, 18, 14])
    xlsx_title_row(so, 1, "RBI Team Sign-off", 5)
    xlsx_header(so, 3, ["#", "Role", "Name", "Cert / Qualification", "Date"])
    for i, role in enumerate(["RBI Facilitator", "Inspection Engineer", "Process / Corrosion Engineer", "Operations Representative", "Asset Owner"], 1):
        xlsx_data_row(so, 3 + i, [i, role, "", "", ""])
    xlsx_brand_footer(so, 10, 5)

    wb.save(f"{OUT}/rbi-worksheet.xlsx")

# ============================================================
# 12) Calibration Certificate Template (DOCX) — ISO/IEC 17025
# ============================================================
def make_calibration_certificate():
    doc = Document()
    docx_setup(doc, "Calibration Certificate — ISO/IEC 17025 §7.8")

    docx_section(doc, "1. Certificate Identification")
    for k, v in [
        ("Certificate Number", "[CAL-YYYY-NNNN]"),
        ("Issue Date", "[YYYY-MM-DD]"),
        ("Page", "1 of __"),
        ("Calibration Laboratory", "[Lab Name]"),
        ("Lab Accreditation", "[ILAC-MRA / A2LA / UKAS / NABL Cert No.]"),
        ("Lab Address", "[Street, City, State, Country]"),
    ]:
        docx_label(doc, k, v)

    docx_section(doc, "2. Customer Information")
    for k, v in [
        ("Customer Name", "[Customer / Owner]"),
        ("Customer Address", "[Full Address]"),
        ("Customer Order / PO No.", "[PO Number]"),
        ("Date Received", "[YYYY-MM-DD]"),
        ("Date of Calibration", "[YYYY-MM-DD]"),
        ("Date Returned", "[YYYY-MM-DD]"),
    ]:
        docx_label(doc, k, v)

    docx_section(doc, "3. Instrument Under Test (IUT)")
    for k, v in [
        ("Instrument Description", "[e.g., Digital UT Flaw Detector]"),
        ("Manufacturer", "[Olympus / Sonatest / GE / Eddyfi]"),
        ("Model Number", "[Model]"),
        ("Serial Number", "[Serial]"),
        ("Asset / ID Tag", "[Internal Tag]"),
        ("Range", "[e.g., 0 - 1000 mm steel]"),
        ("Resolution", "[e.g., 0.01 mm]"),
        ("Condition Received", "[Good / Damaged — describe]"),
        ("Calibration Procedure", "[Lab Procedure / Manufacturer Manual + Issue]"),
        ("Calibration Interval", "[12 months]"),
        ("Next Calibration Due", "[YYYY-MM-DD]"),
    ]:
        docx_label(doc, k, v)

    docx_section(doc, "4. Environmental Conditions")
    for k, v in [
        ("Ambient Temperature", "[deg C]"),
        ("Relative Humidity", "[% RH]"),
        ("Atmospheric Pressure", "[hPa]"),
        ("Lab Location", "[Lab room / Field site]"),
    ]:
        docx_label(doc, k, v)

    docx_section(doc, "5. Reference Standards Used")
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Standard Description", "Make / Model", "Serial No.", "Cal. Cert.", "Traceability"]):
        table.rows[0].cells[i].text = h
    for row in [
        ("[Reference UT Block — IIW V1]", "[Block Spec]", "[Serial]", "[Cert No.]", "[NIST / NPL / PTB]"),
        ("[Step Wedge]", "[Spec]", "[Serial]", "[Cert No.]", "[NIST]"),
        ("[Digital Multimeter]", "[Fluke 87V]", "[Serial]", "[Cert No.]", "[NIST]"),
        ("[Frequency Counter]", "[HP/Keysight]", "[Serial]", "[Cert No.]", "[NIST]"),
    ]:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v

    docx_section(doc, "6. Results — As Found")
    doc.add_paragraph(
        "Measurement results obtained on the instrument upon receipt, before any adjustments. Values reflect "
        "the instrument's condition as presented for calibration."
    )
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Parameter / Point", "Nominal", "Measured", "Error", "Uncertainty (k=2)"]):
        table.rows[0].cells[i].text = h
    for row in [
        ("[Point 1]", "", "", "", ""),
        ("[Point 2]", "", "", "", ""),
        ("[Point 3]", "", "", "", ""),
        ("[Point 4]", "", "", "", ""),
        ("[Point 5]", "", "", "", ""),
    ]:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v

    docx_section(doc, "7. Results — As Left")
    doc.add_paragraph(
        "Measurement results after any adjustments. If no adjustments were performed, as-left values equal "
        "as-found values."
    )
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Parameter / Point", "Nominal", "Measured", "Error", "Uncertainty (k=2)"]):
        table.rows[0].cells[i].text = h
    for row in [
        ("[Point 1]", "", "", "", ""),
        ("[Point 2]", "", "", "", ""),
        ("[Point 3]", "", "", "", ""),
        ("[Point 4]", "", "", "", ""),
        ("[Point 5]", "", "", "", ""),
    ]:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v

    docx_section(doc, "8. Measurement Uncertainty")
    doc.add_paragraph(
        "Uncertainty is reported at a coverage factor of k = 2 corresponding to a confidence level of "
        "approximately 95%. The uncertainty budget includes contributions from the reference standard, "
        "resolution of the IUT, environmental influences, repeatability, and reproducibility, calculated "
        "in accordance with the JCGM 100:2008 (GUM)."
    )
    docx_label(doc, "Expanded Uncertainty (U, k=2)", "[Value + units]")
    docx_label(doc, "Effective Degrees of Freedom", "[v_eff]")

    docx_section(doc, "9. Decision Rule")
    doc.add_paragraph(
        "Statements of conformity, when issued, follow the decision rule agreed with the customer per "
        "ISO/IEC 17025:2017 §7.8.6.1. Where no specific rule is agreed, a simple acceptance rule with no "
        "guard band (binary statement based on measured value vs specification limit, with measurement "
        "uncertainty disclosed in the certificate) shall apply."
    )
    docx_label(doc, "Decision Rule Applied", "[Simple acceptance / Guard-band]")
    docx_label(doc, "Specification Limits Used", "[Manufacturer spec / Customer / Code]")
    docx_label(doc, "Conformity Statement", "[PASS / FAIL — see results tables]")

    docx_section(doc, "10. Statement of Traceability")
    doc.add_paragraph(
        "All measurements reported herein are traceable to the International System of Units (SI) through "
        "the chain of unbroken comparisons to national metrology institutes (NIST, NPL, PTB) or to "
        "intrinsic standards. Reference standards used in this calibration are identified in Section 5."
    )

    docx_section(doc, "11. Opinions & Interpretations")
    doc.add_paragraph(
        "[Optional — per ISO/IEC 17025 §7.8.7, opinions and interpretations may be added by an authorized "
        "individual. Identify the basis for any opinion expressed. Maintain records.]"
    )

    docx_section(doc, "12. Authorization")
    table = doc.add_table(rows=3, cols=3)
    table.style = "Light Grid Accent 1"
    table.rows[0].cells[0].text = "Role"
    table.rows[0].cells[1].text = "Name / Signature"
    table.rows[0].cells[2].text = "Date"
    table.rows[1].cells[0].text = "Calibrated by"
    table.rows[2].cells[0].text = "Reviewed / Approved by (Technical Signatory)"

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("This certificate may not be reproduced, except in full, without the written approval of the laboratory. "
                  "Reported uncertainty includes effects of all known sources at the time of calibration.")
    r.italic = True
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    r.font.size = Pt(9)

    docx_footer(doc)
    doc.save(f"{OUT}/calibration-certificate-template.docx")

# ============================================================
# 13) Welder Qualification Test (WPQR) (DOCX) — ASME Section IX
# ============================================================
def make_welder_qualification():
    doc = Document()
    docx_setup(doc, "Welder Performance Qualification Record (WPQR) — ASME Section IX")

    docx_section(doc, "1. Welder Identification")
    for k, v in [
        ("WPQR No.", "[WPQR-XXX-001]"),
        ("Date of Test", "[YYYY-MM-DD]"),
        ("Employer / Organization", "[Company Name]"),
        ("Welder Name", "[Full Name]"),
        ("Welder ID / Stamp", "[ID]"),
        ("Date of Birth", "[YYYY-MM-DD]"),
        ("Photo on File", "[Yes / No — see Annex]"),
        ("Reference Code", "ASME Section IX (latest edition)"),
        ("Test Witnessed By", "[Witness Name + Cert]"),
        ("Inspection / TPI", "[Authorized Inspector / TPI]"),
    ]:
        docx_label(doc, k, v)

    docx_section(doc, "2. WPS Used for the Test")
    for k, v in [
        ("WPS Reference No.", "[WPS-XXX-001 Rev __]"),
        ("PQR Supporting the WPS", "[PQR No.]"),
        ("Welding Process", "[SMAW / GMAW / FCAW / GTAW / SAW]"),
        ("Type (Manual / Semi-Auto / Auto)", "[Manual / Semi-Auto / Machine / Auto]"),
        ("Power Source", "[DC EP / DC EN / AC]"),
    ]:
        docx_label(doc, k, v)

    docx_section(doc, "3. Essential Variables Tested")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Variable (QW-350)", "Actual Value Used", "Range Qualified", "Notes"]):
        table.rows[0].cells[i].text = h
    for row in [
        ("Process (QW-403)", "", "", ""),
        ("Backing (QW-402.4)", "", "", ""),
        ("Base Metal P-No. (QW-403.16)", "", "", ""),
        ("Filler Metal F-No. (QW-404.15)", "", "", ""),
        ("Position (QW-405.3)", "", "", ""),
        ("Diameter / Thickness Range (QW-403.18)", "", "", ""),
        ("Vertical Progression (QW-405.3)", "", "", ""),
        ("Backing Gas (GTAW/GMAW)", "", "", ""),
        ("Transfer Mode (GMAW)", "", "", ""),
        ("Tungsten Type (GTAW)", "", "", ""),
    ]:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v

    docx_section(doc, "4. Test Coupon")
    for k, v in [
        ("Coupon Description", "[Plate / Pipe — diameter, thickness]"),
        ("Base Metal Specification", "[e.g., SA-516 Gr. 70]"),
        ("Coupon Thickness", "[in / mm]"),
        ("Coupon OD (if pipe)", "[in / mm]"),
        ("Joint Configuration", "[V-Groove / U-Groove / Butt / Fillet]"),
        ("Bevel Angle", "[deg]"),
        ("Root Gap", "[mm / in]"),
        ("Root Face", "[mm / in]"),
        ("Backing Type (Y/N)", "[Yes - material / No]"),
    ]:
        docx_label(doc, k, v)

    docx_section(doc, "5. Filler Metal & Welding Parameters")
    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Pass", "Process", "Filler (Spec/Class)", "Dia. (in)", "Current (A) / V", "Comments"]):
        table.rows[0].cells[i].text = h
    for row in [
        ("Root", "", "", "", "", ""),
        ("Hot pass", "", "", "", "", ""),
        ("Fill 1", "", "", "", "", ""),
        ("Fill 2", "", "", "", "", ""),
        ("Cap", "", "", "", "", ""),
    ]:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v

    docx_section(doc, "6. Test Results")
    docx_section(doc, "6.1 Visual Examination (QW-302.4)", level=2)
    docx_label(doc, "Visual Result", "[Acceptable / Not Acceptable]")
    docx_label(doc, "Visual Examiner / Cert", "[Name, AWS CWI / CSWIP / SNT-TC-1A VT II]")

    docx_section(doc, "6.2 Mechanical Bend Tests (QW-462)", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Specimen", "Type (Side / Face / Root)", "Result (Pass/Fail)", "Comments"]):
        table.rows[0].cells[i].text = h
    for row in [("1", "", "", ""), ("2", "", "", ""), ("3", "", "", ""), ("4", "", "", "")]:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v

    docx_section(doc, "6.3 Radiographic / Ultrasonic Examination (where used)", level=2)
    docx_label(doc, "RT / UT Result", "[Acceptable / Not Acceptable]")
    docx_label(doc, "Examined By / Cert", "[Name + Cert]")
    docx_label(doc, "Acceptance Standard", "[ASME Section IX / Section V]")

    docx_section(doc, "7. Range Qualified")
    doc.add_paragraph(
        "Based on the essential variables and test results above, the welder is qualified for the following:"
    )
    for k, v in [
        ("Processes", "[SMAW / GMAW / etc.]"),
        ("P-No. Range", "[QW-423 — e.g., P1 to P1, P3, P4, P5A]"),
        ("F-No. Range", "[QW-433]"),
        ("Position Range", "[F / H / V / OH or All]"),
        ("Diameter Range", "[per QW-452 / QW-403]"),
        ("Thickness Range (Plate / Pipe)", "[per QW-452]"),
        ("With / Without Backing", "[Backing / No backing]"),
        ("Vertical Progression", "[Uphill / Downhill / Either]"),
    ]:
        docx_label(doc, k, v)

    docx_section(doc, "8. Continuity / Renewal (QW-322)")
    doc.add_paragraph(
        "Welder qualification remains valid for as long as the welder uses the process in production, "
        "with no period in excess of 6 months without using the process. Welder shall provide log evidence "
        "of process use at least every 6 months."
    )
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Date", "Process Used", "Project / Job", "Verified By"]):
        table.rows[0].cells[i].text = h
    for _ in range(6):
        cells = table.add_row().cells
        for c in cells:
            c.text = ""

    docx_section(doc, "9. Certification Statement")
    doc.add_paragraph(
        "We certify that the statements made in this record are correct and that the welds were prepared, "
        "welded, and tested in accordance with the requirements of ASME Section IX (current edition). The "
        "welder identified above is qualified for the welding range specified in Section 7."
    )

    docx_section(doc, "10. Signatures")
    table = doc.add_table(rows=4, cols=4)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Role", "Name", "Signature", "Date"]):
        table.rows[0].cells[i].text = h
    table.rows[1].cells[0].text = "Welder"
    table.rows[2].cells[0].text = "Welding Engineer / Examiner"
    table.rows[3].cells[0].text = "Authorized Inspector / TPI"

    docx_footer(doc)
    doc.save(f"{OUT}/welder-qualification-test-wpqr.docx")

# ============================================================
# 14) NDT Written Practice (DOCX) — ASNT SNT-TC-1A
# ============================================================
def make_written_practice():
    doc = Document()
    docx_setup(doc, "NDT Written Practice — ASNT SNT-TC-1A Compliant")

    docx_section(doc, "1. Introduction & Scope")
    doc.add_paragraph(
        "This Written Practice (WP) establishes the program for qualification and certification of "
        "Nondestructive Testing (NDT) personnel employed by [Employer Name]. It is prepared in accordance "
        "with the Recommended Practice ASNT SNT-TC-1A (current edition) and supplements such requirements "
        "with employer-specific provisions for application within [Employer's] organization."
    )
    doc.add_paragraph(
        "Where customer specifications or governing codes (ASME, AWS, API, ISO, NAS, MIL-STD) impose "
        "stricter requirements than this Written Practice, the more stringent requirement shall govern."
    )
    docx_label(doc, "Employer Name", "[Employer]")
    docx_label(doc, "Effective Date", "[YYYY-MM-DD]")
    docx_label(doc, "WP Document No.", "[WPN-NDT-001 Rev __]")
    docx_label(doc, "Prepared by (NDT Level III)", "[Name, ASNT Level III Cert No.]")
    docx_label(doc, "Approved by", "[Name, Position]")

    docx_section(doc, "2. NDT Methods Covered")
    for m in [
        "Ultrasonic Testing (UT)",
        "Radiographic Testing (RT)",
        "Magnetic Particle Testing (MT)",
        "Liquid Penetrant Testing (PT)",
        "Eddy Current Testing (ET)",
        "Visual Testing (VT)",
        "Phased Array Ultrasonic Testing (PAUT)",
        "Time-of-Flight Diffraction (TOFD)",
        "Leak Testing (LT) — [if applicable]",
        "Acoustic Emission (AE) — [if applicable]",
        "Neutron Radiography (NR) — [if applicable]",
    ]:
        doc.add_paragraph(m, style="List Bullet")

    docx_section(doc, "3. Levels of Qualification")
    doc.add_paragraph(
        "Three levels of personnel qualification are recognized within this Written Practice. Each level "
        "has specific responsibilities and authorities as defined below."
    )
    for level, scope in [
        ("Level I (Trainee)",
         "An individual qualified to perform specific calibrations and specific tests in accordance with "
         "written instructions. Records test results. Cannot independently evaluate, accept, or interpret. "
         "Works under direct supervision of a Level II or Level III."),
        ("Level II",
         "An individual qualified to set up and calibrate equipment, perform tests, interpret and evaluate "
         "results against applicable codes and standards, and document tests. Provides on-the-job training "
         "of Level I personnel. May supervise Level I personnel."),
        ("Level III",
         "An individual qualified to develop, qualify, and approve procedures; establish techniques; "
         "interpret codes and standards; evaluate and consult on the results; train and examine other "
         "personnel; conduct audits. Holds responsibility for the technical content of the Written Practice."),
    ]:
        docx_label(doc, level, scope)

    docx_section(doc, "4. Education, Training, and Experience Requirements")
    doc.add_paragraph(
        "Minimum education, training, and on-the-job experience requirements for each method and level "
        "follow ASNT SNT-TC-1A Table 6.3.1 (or equivalent). Specific hours by method are listed in Annex A."
    )
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Method", "Level", "Min Education", "Training (hr)", "Experience (hr)"]):
        table.rows[0].cells[i].text = h
    for row in [
        ("UT", "I", "HS / equiv.", "40", "210"),
        ("UT", "II", "HS / equiv.", "40 (over Lv I)", "630"),
        ("RT", "I", "HS / equiv.", "40", "210"),
        ("RT", "II", "HS / equiv.", "40 (over Lv I)", "630"),
        ("MT", "I", "HS / equiv.", "12", "130"),
        ("MT", "II", "HS / equiv.", "8 (over Lv I)", "390"),
        ("PT", "I", "HS / equiv.", "4", "70"),
        ("PT", "II", "HS / equiv.", "8 (over Lv I)", "130"),
        ("VT", "I", "HS / equiv.", "8", "70"),
        ("VT", "II", "HS / equiv.", "16 (over Lv I)", "130"),
        ("ET", "II", "HS / equiv.", "40", "1200"),
        ("PAUT", "II", "Level II UT", "80", "1200"),
    ]:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v

    docx_section(doc, "5. Examinations")
    doc.add_paragraph(
        "Each candidate for certification at Level I or Level II shall pass three examinations administered "
        "by the Level III responsible for the method:"
    )
    docx_label(doc, "General Examination", "Principles and theory of the method. Minimum 40 questions for UT / RT, 30 for MT / PT / ET / VT.")
    docx_label(doc, "Specific Examination", "Equipment, operating procedures, codes, and applicable standards used by the employer. Minimum 20 questions.")
    docx_label(doc, "Practical Examination", "Hands-on demonstration of the method on representative test specimens. Includes equipment setup, calibration, indication detection, sizing, evaluation, and reporting.")
    doc.add_paragraph(
        "Pass mark: 70% on each examination, with composite (combined) score of 80% for Level II "
        "(per SNT-TC-1A and many customer specifications). Records of all examinations shall be retained."
    )

    docx_section(doc, "6. Vision Requirements")
    for item in [
        "Near Vision: Jaeger No. 2 (or equivalent — Times Roman 4.5) at minimum 12 in. natural / corrected, "
        "in at least one eye. Examined annually.",
        "Color Vision: Ability to distinguish and differentiate among colors used in the method (e.g., "
        "Ishihara plates). Examined upon employment and annually.",
        "Records of vision examinations shall be maintained.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    docx_section(doc, "7. Certification")
    doc.add_paragraph(
        "Upon successful completion of the training, experience, and examinations, the candidate is "
        "certified by the employer's Level III for the specific method and level. Certification records "
        "shall include: certificate number, name and ID of the certified individual, method, level, dates "
        "of certification and expiration, examination scores, training records, experience records, vision "
        "test records, and signature of the certifying Level III."
    )
    docx_label(doc, "Certification Period", "5 years (Level III may shorten if required by customer / risk)")
    docx_label(doc, "Renewal", "Annual review by Level III; full recertification per Section 8")

    docx_section(doc, "8. Recertification")
    doc.add_paragraph(
        "Recertification may be conducted by any of the following options at the discretion of the Level III "
        "(per SNT-TC-1A Section 9):"
    )
    for item in [
        "(a) Evidence of continuing satisfactory performance — Level III's review of work performed during the certification period.",
        "(b) Re-examination — full general, specific, and practical examinations.",
        "(c) Documented evidence of training and / or experience equivalent to the original certification.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    docx_section(doc, "9. Suspension / Revocation of Certification")
    doc.add_paragraph(
        "Certification shall be suspended or revoked by the Level III on evidence of: misconduct related to "
        "NDT activity; failure to comply with this Written Practice or related procedures; loss of vision "
        "requirements; or a period of more than 12 consecutive months without active practice in the certified "
        "method. Reinstatement requires re-examination."
    )

    docx_section(doc, "10. Records & Retention")
    doc.add_paragraph(
        "All records required by this Written Practice (training, experience, examination results, vision "
        "tests, certification certificates) shall be retained for the entire period of employment plus "
        "[5 years] post-departure, in a controlled environment with retrieval traceability per ISO 9001 §7.5."
    )

    docx_section(doc, "11. Audit & Review")
    doc.add_paragraph(
        "This Written Practice shall be reviewed at least annually by the responsible Level III and updated "
        "as needed to reflect changes in SNT-TC-1A, applicable codes, or employer organization. The review "
        "shall be documented with the reviewer's signature and date."
    )

    docx_section(doc, "12. Annex A — Method Hour Requirements Summary")
    doc.add_paragraph("Referenced in Section 4. Hours listed are SNT-TC-1A 2024 minima.")

    docx_section(doc, "13. Annex B — Examination Question Bank Reference")
    doc.add_paragraph("Question banks for each method are maintained by the responsible Level III in a "
                      "controlled environment, separate from training materials. Bank contents are "
                      "referenced by question ID, not by content.")

    docx_section(doc, "14. Revision History")
    table = doc.add_table(rows=2, cols=4)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Rev", "Date", "Description", "Approved by"]):
        table.rows[0].cells[i].text = h
    for i, v in enumerate(["0", "[YYYY-MM-DD]", "Initial issue", "[Level III Name]"]):
        table.rows[1].cells[i].text = v

    docx_footer(doc)
    doc.save(f"{OUT}/ndt-written-practice-template.docx")

# ============================================================
# 15) Inspection Test Plan (ITP) (XLSX)
# ============================================================
def make_itp():
    wb = Workbook()
    ws = wb.active
    ws.title = "ITP"
    xlsx_set_widths(ws, [6, 14, 36, 16, 16, 26, 14, 14, 14, 14, 14])

    xlsx_title_row(ws, 1, "Inspection & Test Plan (ITP) — Construction QA/QC", 11)
    ws.cell(row=2, column=1, value="Project:").font = BOLD
    ws.cell(row=2, column=5, value="Client:").font = BOLD
    ws.cell(row=2, column=9, value="ITP No.:").font = BOLD
    ws.cell(row=3, column=1, value="Discipline:").font = BOLD
    ws.cell(row=3, column=5, value="Revision:").font = BOLD
    ws.cell(row=3, column=9, value="Date:").font = BOLD
    ws.cell(row=4, column=1, value="Reference Specifications:").font = BOLD
    ws.cell(row=4, column=5, value="Codes:").font = BOLD

    # Legend
    xlsx_section(ws, 6, "Legend: H = Hold Point | W = Witness | R = Review of records | M = Monitor | S = Surveillance | I = Information only", 11, color=BRAND_LIGHT)

    # Headers
    xlsx_header(ws, 8, ["#", "ITP Code", "Activity / Inspection Item", "Method", "Acceptance Criteria", "Reference Document", "Frequency", "Contractor", "Sub-Contractor", "Client / TPI", "Records"])

    activities = [
        ("01", "ITP-001", "Material receiving inspection (MTR review)", "Document review + visual", "Per PO / Spec. + EN 10204 3.1", "ASME II / EN 10204", "Each heat / lot", "M", "M", "R", "MRR + MTR"),
        ("02", "ITP-002", "Welding Procedure Specification (WPS) review", "Doc review", "ASME Section IX qualified", "ASME IX / WPS", "Each WPS", "R", "—", "R", "WPS + PQR"),
        ("03", "ITP-003", "Welder qualification verification", "Doc review", "ASME Section IX QW-300", "ASME IX", "Each welder", "M", "—", "R", "WPQ records"),
        ("04", "ITP-004", "Joint fit-up inspection", "Visual + measurement", "Per drawing + WPS tolerances", "Drawing / WPS", "Each joint", "M", "M", "W", "Fit-up report"),
        ("05", "ITP-005", "Pre-heat verification", "Temperature measurement", "Min per WPS / Code", "WPS / ASME B31.3", "Each joint when applicable", "M", "M", "S", "Pre-heat log"),
        ("06", "ITP-006", "Visual inspection of welds (VT)", "VT per ASME V Article 9", "Per code (ASME Section VIII / B31.3)", "ASME V / VIII / B31.3", "100% of welds", "M", "M", "W/R", "VT report"),
        ("07", "ITP-007", "Magnetic Particle Testing (MT)", "Yoke / prod / continuous", "ASME VIII Appendix 6", "ASME V Article 7", "Per project spec.", "M", "M", "R", "MT report"),
        ("08", "ITP-008", "Liquid Penetrant Testing (PT)", "Solvent removable / water-washable", "ASME VIII Appendix 8", "ASME V Article 6", "Per project spec.", "M", "M", "R", "PT report"),
        ("09", "ITP-009", "Ultrasonic Testing (UT)", "Manual / PAUT", "ASME VIII Appendix 12 / 4", "ASME V Article 4 / 5", "Per project spec.", "M", "M", "W/R", "UT report"),
        ("10", "ITP-010", "Radiographic Testing (RT)", "Gamma / X-ray", "ASME VIII Appendix 4", "ASME V Article 2", "Per project spec.", "M", "M", "W/R", "RT film/DR + report"),
        ("11", "ITP-011", "Post-Weld Heat Treatment (PWHT)", "Heat treatment", "Per WPS / Code", "ASME VIII UCS-56", "Each cycle", "H", "M", "W", "PWHT chart + record"),
        ("12", "ITP-012", "Hardness Testing", "Brinell / Vickers", "ASME B31.3 Table 331.1.1", "ASME B31.3", "Per spec.", "M", "M", "R", "Hardness report"),
        ("13", "ITP-013", "Positive Material Identification (PMI)", "OES / XRF", "Per material spec.", "API 578", "100% / sampling", "M", "M", "R", "PMI report"),
        ("14", "ITP-014", "Dimensional / final visual inspection", "VT + measurement", "Per drawing tolerances", "Drawing", "100% of assemblies", "M", "M", "W", "Dimensional report"),
        ("15", "ITP-015", "Pressure Test (hydrotest / pneumatic)", "Hydro / pneumatic", "Per code", "ASME VIII UG-99 / B31.3 §345", "Each system", "H", "M", "W", "Pressure-test record"),
        ("16", "ITP-016", "Surface preparation & painting", "VT + thickness", "Per spec.", "SSPC / NACE", "Each coat", "M", "M", "S/R", "Coating report"),
        ("17", "ITP-017", "Final documentation review", "Doc review", "Per project QA plan", "Project QA Plan", "Project close", "R", "—", "R", "Final dossier / MRB"),
    ]
    row = 9
    for a in activities:
        xlsx_data_row(ws, row, list(a))
        row += 1

    # Sign-off
    row += 1
    xlsx_section(ws, row, "ITP Approval", 11, color=BRAND_LIGHT)
    row += 1
    xlsx_header(ws, row, ["Role", "Name", "Position / Cert", "Signature", "Date", "Comments", "—", "—", "—", "—", "—"])
    for r in ["Contractor QC Manager", "Welding Engineer", "Sub-Contractor QC", "Client Project Manager", "TPI / Third Party"]:
        row += 1
        xlsx_data_row(ws, row, [r, "", "", "", "", "", "", "", "", "", ""])

    row += 2
    xlsx_brand_footer(ws, row, 11)
    wb.save(f"{OUT}/inspection-test-plan-itp.xlsx")

# ============================================================
# 16) Audit / NCR / CAPA Finding Tracker (XLSX)
# ============================================================
def make_audit_tracker():
    wb = Workbook()
    ws = wb.active
    ws.title = "Findings Tracker"
    xlsx_set_widths(ws, [6, 14, 14, 14, 14, 30, 30, 22, 14, 14, 14, 14, 22])

    xlsx_title_row(ws, 1, "Audit Findings / NCR / CAPA Tracker", 13)
    ws.cell(row=2, column=1, value="Organization:").font = BOLD
    ws.cell(row=2, column=6, value="Reporting Period:").font = BOLD
    ws.cell(row=2, column=10, value="Owner:").font = BOLD
    ws.cell(row=3, column=1, value="Standards:").font = BOLD
    ws.cell(row=3, column=6, value="ISO 9001:2015 | ISO 17025 | ISO 45001 | API Q1 / Q2").font = Font(name="Calibri", size=10, italic=True)

    # Legend
    xlsx_section(ws, 5, "Severity: 1 Critical | 2 Major | 3 Minor | 4 Observation | 5 Opportunity for Improvement", 13, color=BRAND_LIGHT)

    # Headers
    xlsx_header(ws, 7, ["#", "Finding No.", "Date Raised", "Source", "Category", "Description of Finding", "Root Cause", "Corrective / Preventive Action", "Severity", "Owner", "Target Date", "Status", "Effectiveness Review"])

    sample_categories = ["NDT Procedure", "Personnel Cert.", "Equipment Calibration", "Reporting / Records", "Customer Complaint", "Safety", "Document Control", "Training", "Supplier", "Internal Audit", "External Audit", "Management Review"]

    # 20 empty rows for users to fill, with pre-filled categories on first 12
    for i in range(1, 21):
        cat = sample_categories[(i - 1) % len(sample_categories)] if i <= 12 else ""
        xlsx_data_row(ws, 7 + i, [i, f"NCR-{i:04d}", "", "", cat, "", "", "", "", "", "", "Open", ""])

    # Dashboard counts
    row = 7 + 20 + 2
    xlsx_section(ws, row, "Status Dashboard", 13, color=BRAND, dark=True)
    row += 1
    xlsx_header(ws, row, ["Status", "Count", "—", "Severity", "Count", "—", "Category", "Count", "—", "Source", "Count", "—", "Comments"])
    statuses = ["Open", "In Progress", "Pending Verification", "Closed - Effective", "Closed - Not Effective", "Re-opened"]
    severities = ["1 Critical", "2 Major", "3 Minor", "4 Observation", "5 OFI"]
    for i, s in enumerate(statuses):
        row += 1
        xlsx_data_row(ws, row, [s, "", "", severities[i] if i < len(severities) else "", "", "", sample_categories[i] if i < len(sample_categories) else "", "", "", "", "", "", ""])

    # Sign-off
    row += 2
    xlsx_section(ws, row, "Tracker Sign-off / Management Review", 13, color=BRAND_LIGHT)
    row += 1
    xlsx_header(ws, row, ["Role", "Name", "Position", "Signature", "Date", "Comments", "—", "—", "—", "—", "—", "—", "—"])
    for role in ["QA Manager", "Top Management Representative", "Process Owner", "Internal Auditor"]:
        row += 1
        xlsx_data_row(ws, row, [role, "", "", "", "", "", "", "", "", "", "", "", ""])

    row += 2
    xlsx_brand_footer(ws, row, 13)
    wb.save(f"{OUT}/audit-finding-tracker.xlsx")

# ============================================================
# Run
# ============================================================
if __name__ == "__main__":
    make_inspection_checklist()
    make_api653_template()
    make_safety_checklist()
    make_training_matrix()
    make_procedure_template()
    make_level_iii_study_guide()
    # New (Sprint 2)
    make_api510_report()
    make_api570_record()
    make_dpr()
    make_pwht_record()
    make_rbi_worksheet()
    make_calibration_certificate()
    make_welder_qualification()
    make_written_practice()
    make_itp()
    make_audit_tracker()
    print("Templates generated in", OUT)
    for f in sorted(os.listdir(OUT)):
        size = os.path.getsize(f"{OUT}/{f}")
        print(f"  {f}  ({size:,} bytes)")
